#!/usr/bin/env python3
"""Build a DOCX manuscript draft with main figures and tables inserted."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
MANUSCRIPT = ROOT / "manuscript/manuscript_full_zh_v5.md"
TABLES = ROOT / "manuscript/tables_draft_zh_v5.md"
LEGENDS = ROOT / "manuscript/figure_legends_zh_v5.md"
OUT = ROOT / "submission_package_v5/manuscript/manuscript_full_zh_v5.docx"

FIGURES = {
    "Fig. 1": ROOT / "figures/manuscript_main/figure1_benchmark_design.png",
    "Fig. 2": ROOT / "figures/manuscript_main/figure2_baseline_decay.png",
    "Fig. 3": ROOT / "figures/manuscript_main/figure3_topk_response_gene_recovery.png",
    "Fig. 4": ROOT / "figures/manuscript_main/figure4_chemical_similarity_audit.png",
    "Fig. 5": ROOT / "figures/manuscript_main/figure6_ranking_instability.png",
    "Fig. 6": ROOT / "figures/manuscript_main/figure5_svd_ridge_stress_test.png",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.size = Pt(8)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Heading 1", 15, "1f3f41"),
        ("Heading 2", 12, "2f6f73"),
        ("Heading 3", 10.5, "2f3437"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def clean_inline(text: str) -> str:
    return text.replace("`", "")


def parse_figure_legends() -> dict[str, str]:
    text = LEGENDS.read_text(encoding="utf-8")
    legends = {}
    for block in re.split(r"\n## Figure ", text):
        if not block.strip() or block.startswith("#"):
            continue
        first, *rest = block.strip().splitlines()
        num = first.split(".", 1)[0].strip()
        body = " ".join(line.strip() for line in rest if line.strip() and not line.startswith("对应文件"))
        legends[f"Fig. {num}"] = f"Figure {num}. {body}"
    return legends


def parse_markdown_tables() -> dict[str, list[list[str]]]:
    text = TABLES.read_text(encoding="utf-8")
    tables = {}
    current = None
    rows: list[list[str]] = []
    for line in text.splitlines():
        if line.startswith("## Table "):
            if current and rows:
                tables[current] = rows
            current = "Table " + line.split("Table ", 1)[1].split(".", 1)[0]
            rows = []
            continue
        if current and line.startswith("|") and "---" not in line:
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(cells)
        elif current and rows and line.strip() == "":
            tables[current] = rows
            current = None
            rows = []
    if current and rows:
        tables[current] = rows
    return tables


def add_table(doc: Document, key: str, rows: list[list[str]]) -> None:
    if key in {"Table 5", "Table 6", "Table 8"}:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    p.add_run(key).bold = True
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            set_cell_text(table.cell(i, j), text, bold=(i == 0))
            if i == 0:
                set_cell_shading(table.cell(i, j), "E9EEF3")
    doc.add_paragraph()


def add_figure(doc: Document, key: str, legend: str) -> None:
    image = FIGURES[key]
    if key != "Fig. 5":
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(6.25 if key == "Fig. 5" else 6.7)
    p.add_run().add_picture(str(image), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap.add_run(legend)
    run.font.size = Pt(8)
    run.italic = True


def add_markdown_paragraph(doc: Document, line: str) -> None:
    if line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean_inline(line[2:]))
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("1f3f41")
    elif line.startswith("## "):
        doc.add_heading(clean_inline(line[3:]), level=1)
    elif line.startswith("### "):
        doc.add_heading(clean_inline(line[4:]), level=2)
    elif re.match(r"^\d+\. ", line):
        doc.add_paragraph(clean_inline(line), style=None)
    elif line.strip():
        doc.add_paragraph(clean_inline(line))


def main() -> None:
    doc = Document()
    style_document(doc)
    legends = parse_figure_legends()
    tables = parse_markdown_tables()
    inserted_figures: set[str] = set()
    inserted_tables: set[str] = set()

    for raw in MANUSCRIPT.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        add_markdown_paragraph(doc, line)
        for key in FIGURES:
            if key in line and key not in inserted_figures:
                add_figure(doc, key, legends.get(key, key))
                inserted_figures.add(key)
        for i in range(1, 9):
            key = f"Table {i}"
            if key in line and key in tables and key not in inserted_tables:
                add_table(doc, key, tables[key])
                inserted_tables.add(key)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
